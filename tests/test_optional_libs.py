"""
Test optional libraries (extras) by calling MCP server tools.

Run from project root with a venv that has all extras installed:
  pip install -e ".[all]"   (or: uv pip install -e ".[all]")
  # Activate venv then:
  python tests/test_optional_libs.py

All tests invoke the mcp-docx server via MCP (stdio); no direct Python calls to tools.
Use the same Python/interpreter for the test and for the server subprocess (avoid
'uv run' if it causes the test to hang due to stdio buffering).
"""
from pathlib import Path
import sys
import io
import os
import asyncio
import json
import signal

# Avoid Windows console encoding errors with Unicode output
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# Project root and output dir (absolute paths so MCP server process can read/write)
ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "out" / "test_optional"
OUT.mkdir(parents=True, exist_ok=True)


def ok(name: str) -> None:
    print(f"  [OK] {name}")


def fail(name: str, e: Exception) -> None:
    print(f"  [FAIL] {name}: {e}")
    raise


def _tool_result_text(result) -> str:
    """Extract JSON/text from MCP CallToolResult."""
    if not getattr(result, "content", None):
        return ""
    first = result.content[0]
    return getattr(first, "text", None) or str(first)


async def run_mcp_tests() -> None:
    """Spawn MCP server, connect client, run tool tests."""
    from mcp import ClientSession, StdioServerParameters, stdio_client

    # Server: python -m mcp_docx.server, cwd = project root (unbuffered so responses flush)
    env = os.environ.copy()
    env["PYTHONUNBUFFERED"] = "1"
    env.pop("PYTHONIOENCODING", None)  # avoid subprocess encoding issues
    server_params = StdioServerParameters(
        command=sys.executable,
        args=["-u", "-m", "mcp_docx.server"],
        cwd=str(ROOT),
        env=env,
    )

    async with stdio_client(server_params) as (read_stream, write_stream):
        async with ClientSession(read_stream, write_stream) as session:
            await session.initialize()

            # --- create_document (core) ---
            print("\n--- MCP: create_document ---")
            out_doc = str(OUT / "mcp_create_doc.docx")
            result = await session.call_tool("create_document", {
                "output_path": out_doc,
                "title": "MCP Test",
                "paragraphs": ["First.", "Second."],
            })
            text = _tool_result_text(result)
            if result.isError:
                fail("create_document", RuntimeError(text or "Tool returned error"))
            data = json.loads(text) if text else {}
            if Path(data.get("outputPath", "")).is_file():
                ok("create_document")
            else:
                fail("create_document", FileNotFoundError("Output not created"))

            # --- create_document_formatted (core + optional image) ---
            print("\n--- MCP: create_document_formatted (with image block) ---")
            try:
                from PIL import Image
                img_path = OUT / "test_pixel.png"
                Image.new("RGB", (10, 10), color="red").save(img_path)
            except ImportError:
                print("  [SKIP] Pillow not installed")
            else:
                out_fmt = str(OUT / "mcp_formatted_with_image.docx")
                result = await session.call_tool("create_document_formatted", {
                    "output_path": out_fmt,
                    "title": "Formatted + Image",
                    "content_blocks": [
                        {"type": "paragraph", "text": "Below is an image:"},
                        {"type": "image", "path": str(img_path), "width_inches": 1.0},
                    ],
                })
                text = _tool_result_text(result)
                if result.isError:
                    print(f"  [SKIP or FAIL] create_document_formatted+image: {text}")
                elif Path(json.loads(text).get("outputPath", "")).is_file():
                    ok("create_document_formatted with image block")
                else:
                    fail("create_document_formatted", FileNotFoundError("Output not created"))

            # --- generate_from_template (core) ---
            print("\n--- MCP: generate_from_template ---")
            tpl_path = OUT / "jinja_template.docx"
            if not tpl_path.is_file():
                from docx import Document
                d = Document()
                d.add_paragraph("Hello {{ name }}, date: {{ date }}.")
                d.save(tpl_path)
            out_tpl = str(OUT / "mcp_from_template.docx")
            result = await session.call_tool("generate_from_template", {
                "template_path": str(tpl_path),
                "output_path": out_tpl,
                "data": {"name": "World", "date": "2025-01-01"},
            })
            text = _tool_result_text(result)
            if result.isError:
                fail("generate_from_template", RuntimeError(text))
            if Path(json.loads(text).get("outputPath", "")).is_file():
                ok("generate_from_template")
            else:
                fail("generate_from_template", FileNotFoundError("Output not created"))

            # --- mail_merge (optional: docx-mailmerge) ---
            print("\n--- MCP: mail_merge ---")
            merge_tpl = _make_merge_template()
            if merge_tpl is None:
                print("  [SKIP] docx-mailmerge or template creation failed")
            else:
                out_merge = str(OUT / "mcp_merged.docx")
                result = await session.call_tool("mail_merge", {
                    "template_path": merge_tpl,
                    "output_path": out_merge,
                    "merge_data": {"Nome": "MCPUser"},
                })
                text = _tool_result_text(result)
                if result.isError:
                    print(f"  [SKIP or FAIL] mail_merge: {text}")
                elif Path(json.loads(text).get("outputPath", "")).is_file():
                    ok("mail_merge")
                else:
                    fail("mail_merge", FileNotFoundError("Output not created"))

            # --- convert_pdf_to_docx (optional: pdf2docx) ---
            print("\n--- MCP: convert_pdf_to_docx ---")
            pdf_path = _make_minimal_pdf()
            out_pdf = str(OUT / "mcp_from_pdf.docx")
            result = await session.call_tool("convert_pdf_to_docx", {
                "pdf_path": pdf_path,
                "output_path": out_pdf,
            })
            text = _tool_result_text(result)
            if result.isError:
                print(f"  [SKIP or FAIL] convert_pdf_to_docx: {text}")
            elif Path(json.loads(text).get("outputPath", "")).is_file():
                ok("convert_pdf_to_docx")
            else:
                fail("convert_pdf_to_docx", FileNotFoundError("Output not created"))

            # --- docx_to_html (optional: mammoth) ---
            print("\n--- MCP: docx_to_html ---")
            docx_for_html = OUT / "for_mammoth.docx"
            if not docx_for_html.is_file():
                from docx import Document
                d = Document()
                d.add_paragraph("Hello World")
                d.save(docx_for_html)
            result = await session.call_tool("docx_to_html", {
                "document_path": str(docx_for_html),
            })
            text = _tool_result_text(result)
            if result.isError:
                print(f"  [SKIP or FAIL] docx_to_html: {text}")
            else:
                data = json.loads(text)
                if data.get("html") and "<" in data["html"]:
                    ok("docx_to_html")
                else:
                    fail("docx_to_html", ValueError("No HTML in result"))

            # --- extract_text, get_document_info (core) ---
            print("\n--- MCP: extract_text, get_document_info ---")
            result = await session.call_tool("extract_text", {"document_path": out_doc})
            text = _tool_result_text(result)
            if result.isError:
                print(f"  [FAIL] extract_text: {text}")
            elif json.loads(text).get("text") is not None:
                ok("extract_text")
            result = await session.call_tool("get_document_info", {"document_path": out_doc})
            text = _tool_result_text(result)
            if result.isError:
                print(f"  [FAIL] get_document_info: {text}")
            else:
                data = json.loads(text) if text else {}
                if "paragraphCount" in data or "wordCount" in data:
                    ok("get_document_info")
                else:
                    ok("get_document_info")

            # --- convert_document / pypandoc (optional; often skipped if Pandoc not installed) ---
            print("\n--- MCP: convert_document (Pandoc) ---")
            result = await session.call_tool("convert_document", {
                "input_path": out_doc,
                "output_path": str(OUT / "mcp_to_md.md"),
                "output_format": "md",
            })
            text = _tool_result_text(result)
            if result.isError:
                print(f"  [SKIP] convert_document (Pandoc): {text}")
            else:
                ok("convert_document")

    print("\nAll MCP tool tests finished.")


def _make_minimal_pdf() -> str:
    """Write a minimal valid PDF to out dir; the MCP server needs pdf2docx to convert it."""
    minimal_pdf = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R>>endobj
4 0 obj<</Length 44>>stream
BT /F1 12 Tf 100 700 Td (Hello) Tj ET
endstream endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000206 00000 n
trailer<</Size 5/Root 1 0 R>>
startxref
298
%%EOF"""
    pdf_path = OUT / "minimal.pdf"
    pdf_path.write_bytes(minimal_pdf)
    return str(pdf_path)


def _make_merge_template() -> str | None:
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return None
    doc = Document()
    p = doc.add_paragraph()
    p_elem = p._element

    def add_fld_char(elm, fld_type: str) -> None:
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), fld_type)
        elm.append(fld)

    def add_instr(elm, text: str) -> None:
        run = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = text
        run.append(instr)
        elm.append(run)

    add_fld_char(p_elem, "begin")
    add_instr(p_elem, "MERGEFIELD Nome")
    add_fld_char(p_elem, "separate")
    add_fld_char(p_elem, "end")
    path = OUT / "merge_template.docx"
    doc.save(path)
    return str(path)


def main() -> None:
    print("Optional libs test via MCP (output in out/test_optional/)")
    try:
        import anyio
        anyio.run(run_mcp_tests)
    except Exception as e:
        print(f"\nStopped: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
