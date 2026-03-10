"""Shared helpers for .docx operations."""

import unicodedata
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def normalize_text(s: str) -> str:
    """Normalize text to Unicode NFC for consistent accented characters."""
    return unicodedata.normalize("NFC", s)


def ensure_dir(path: Path) -> None:
    """Ensure parent directory exists."""
    path.parent.mkdir(parents=True, exist_ok=True)


def set_document_default_font(doc: Document, font_name: str = "Calibri", size_pt: float = 11) -> None:
    """Set the Normal style default font (Latin and East Asian) for consistent body text."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    try:
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def create_minimal_document(title: str, paragraphs: list[str]) -> Document:
    """Create a new Document with title and paragraphs."""
    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(normalize_text(title))
    run.bold = True
    run.font.size = Pt(24)
    for text in paragraphs:
        doc.add_paragraph(normalize_text(text))
    return doc


def create_document_formatted(
    content_blocks: list[dict[str, Any]],
    title: str | None = None,
    subtitle: str | None = None,
    font_name: str = "Calibri",
    font_size_pt: float = 11,
) -> Document:
    """
    Create a Document with optional title/subtitle and rich content blocks.
    content_blocks: list of dicts, each with:
      - type "heading": level (1 or 2), text
      - type "paragraph": text
      - type "list_bullet": items (list of strings)
    """
    doc = Document()
    set_document_default_font(doc, font_name, font_size_pt)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(normalize_text(title))
        run.bold = True
        run.font.size = Pt(18)
        p.style = doc.styles["Title"]
    if subtitle:
        p = doc.add_paragraph(normalize_text(subtitle))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    for block in content_blocks:
        btype = block.get("type", "paragraph")
        if btype == "heading":
            level = min(max(int(block.get("level", 1)), 1), 2)
            text = block.get("text", "")
            doc.add_heading(normalize_text(text), level=level)
        elif btype == "list_bullet":
            items = block.get("items") or []
            for item in items:
                doc.add_paragraph(normalize_text(str(item)), style="List Bullet")
        else:
            text = block.get("text", "")
            doc.add_paragraph(normalize_text(text))
    return doc
